#!/usr/bin/env python3
"""
Conversion du rapport de qualité en document Word
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def setup_professional_styles(doc):
    """Configure des styles professionnels Word"""

    # Style titre principal
    title_style = doc.styles['Title']
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.name = 'Arial'
    title_style.font.color.rgb = RGBColor(31, 119, 180)

    # Style sous-titre
    subtitle_style = doc.styles['Subtitle']
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.italic = True
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.color.rgb = RGBColor(89, 89, 89)

    # Styles de titre avec niveaux
    for level in range(1, 4):
        style_name = f'Heading{level}Custom'
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        if level == 1:
            style.font.size = Pt(20)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            style.font.size = Pt(16)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(6)

        style.font.bold = True
        style.font.name = 'Arial'

    # Style pour le code
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(64, 64, 64)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Style pour les métriques
    metric_style = doc.styles.add_style('MetricStyle', WD_STYLE_TYPE.PARAGRAPH)
    metric_style.font.size = Pt(11)
    metric_style.font.name = 'Arial'
    metric_style.paragraph_format.left_indent = Inches(0.2)

    # Style pour les scores
    score_style = doc.styles.add_style('ScoreStyle', WD_STYLE_TYPE.PARAGRAPH)
    score_style.font.size = Pt(12)
    score_style.font.bold = True
    score_style.font.name = 'Arial'

def convert_markdown_to_word(markdown_content, doc):
    """Convertit le contenu Markdown en Word"""

    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Titres
        if line.startswith('# '):
            title = doc.add_paragraph(line[2:], style='Title')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading1Custom')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading2Custom')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Heading3Custom')

        # Lignes vides
        elif line == '':
            if i > 0 and lines[i-1].strip() != '':
                doc.add_paragraph('')

        # Listes
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')

        # Texte normal avec formatage
        elif line:
            p = doc.add_paragraph()

            # Traitement du formatage markdown dans le texte
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)

            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

        i += 1

def create_quality_report_word():
    """Crée le rapport de qualité en Word"""

    print("📄 Génération du rapport Word en cours...")

    # Créer le document
    doc = Document()

    # Styles professionnels
    setup_professional_styles(doc)

    # Page de garde
    title = doc.add_paragraph("📊 Rapport de Qualité du Code", style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Consultator - Application de Gestion", style='Subtitle')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.add_run("Analyse complète de la qualité du code\n").bold = False
    info.add_run("Métriques, sécurité et recommandations d'amélioration\n").bold = False
    info.add_run("Généré automatiquement").bold = False
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    version_info = doc.add_paragraph()
    version_info.add_run("Version 1.0.0\n").bold = True
    version_info.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}\n").bold = False
    version_info.add_run("Branche: Master").bold = False
    version_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Saut de page
    doc.add_page_break()

    # Vérifier si le rapport markdown existe
    markdown_path = "docs/quality_report.md"
    if os.path.exists(markdown_path):
        print(f"📖 Lecture du rapport markdown: {markdown_path}")

        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Convertir en Word
        convert_markdown_to_word(markdown_content, doc)

    else:
        print("⚠️ Rapport markdown non trouvé, génération d'un rapport de base...")

        # Rapport de base si le markdown n'existe pas
        doc.add_paragraph("Rapport de Qualité du Code - Consultator", style='Heading1Custom')

        doc.add_paragraph("Ce rapport contient l'analyse de la qualité du code de l'application Consultator.", style='Normal')

        doc.add_paragraph("Métriques Clés:", style='Heading2Custom')

        metrics = [
            "• Structure du projet",
            "• Métriques de code",
            "• Analyse de sécurité",
            "• Problèmes de qualité",
            "• Recommandations d'amélioration"
        ]

        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')

    # Sauvegarder le document
    output_path = "docs/Rapport_Qualite_Code_Consultator.docx"
    doc.save(output_path)

    print(f"✅ Rapport Word créé: {output_path}")
    print(f"📊 Taille: {os.path.getsize(output_path)} octets")

    return output_path

if __name__ == "__main__":
    from datetime import datetime
    create_quality_report_word()
