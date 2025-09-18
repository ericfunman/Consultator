#!/usr/bin/env python3
"""
Script pour convertir la documentation Sphinx RST en document Word
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml


def setup_document_styles(doc):
    """Configure les styles du document Word"""

    # Style pour les titres principaux
    title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.name = "Arial"
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Style pour les titres de niveau 1
    h1_style = doc.styles.add_style("Heading1Custom", WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.name = "Arial"
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Style pour les titres de niveau 2
    h2_style = doc.styles.add_style("Heading2Custom", WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.name = "Arial"
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Style pour les titres de niveau 3
    h3_style = doc.styles.add_style("Heading3Custom", WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.name = "Arial"
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # Style pour le code
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.25)

    # Style pour les listes
    list_style = doc.styles.add_style("CustomList", WD_STYLE_TYPE.PARAGRAPH)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)


def convert_rst_to_docx(rst_content, doc):
    """Convertit le contenu RST en éléments Word"""

    lines = rst_content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Détecter les titres
        if i + 1 < len(lines) and lines[i + 1].startswith(("=", "-", "~", "^", '"')):
            title_text = line.strip()
            underline = lines[i + 1][0]

            if underline == "=":
                # Titre principal
                p = doc.add_paragraph(title_text, style="CustomTitle")
            elif underline == "-":
                # Titre niveau 1
                p = doc.add_paragraph(title_text, style="Heading1Custom")
            elif underline == "~":
                # Titre niveau 2
                p = doc.add_paragraph(title_text, style="Heading2Custom")
            elif underline == "^":
                # Titre niveau 3
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            else:
                # Titre niveau 4 ou plus
                p = doc.add_paragraph(title_text, style="Heading3Custom")

            i += 2  # Sauter la ligne de soulignement
            continue

        # Détecter les blocs de code
        if (
            line.startswith(".. code-block::")
            or line.startswith("   ")
            or (
                line.startswith("   ")
                and i > 0
                and lines[i - 1].startswith(".. code-block::")
            )
        ):
            code_lines = []
            j = i

            # Collecter toutes les lignes de code
            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith("   ") or (
                    j == i and line.startswith(".. code-block::")
                ):
                    if line.startswith(".. code-block::"):
                        # Sauter la directive
                        j += 1
                        continue
                    code_lines.append(current_line[3:])  # Retirer l'indentation
                    j += 1
                else:
                    break

            if code_lines:
                # Ajouter le bloc de code
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph(code_text, style="CodeBlock")
                i = j
                continue

        # Détecter les listes
        if line.startswith(("- ", "* ", "+ ", "1. ", "2. ", "3. ")):
            p = doc.add_paragraph(line[2:], style="CustomList")
            i += 1
            continue

        # Détecter les directives à ignorer
        if line.startswith(".. ") or line.strip() == "":
            i += 1
            continue

        # Texte normal
        if line.strip():
            p = doc.add_paragraph(line)
            # Détecter les éléments en gras ou italique
            # Version simplifiée - on pourrait améliorer avec regex
            if "**" in line:
                # Traitement basique du gras
                pass

        i += 1


def create_word_documentation():
    """Crée la documentation complète en Word"""

    # Fichiers à traiter dans l'ordre
    files_order = [
        "index.rst",
        "installation.rst",
        "quickstart.rst",
        "features.rst",
        "development.rst",
        "api.rst",
        "tutorials.rst",
    ]

    # Créer le document
    doc = Document()

    # Configuration des styles
    setup_document_styles(doc)

    # Titre du document
    title = doc.add_paragraph("Documentation Consultator", style="CustomTitle")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Informations générales
    doc.add_paragraph("Application de gestion d'une practice data de consultants")
    doc.add_paragraph("Version 1.0.0 - Générée automatiquement depuis Sphinx")
    doc.add_paragraph("Date: Septembre 2025")
    doc.add_paragraph("")

    # Traiter chaque fichier
    docs_dir = Path(__file__).parent

    for filename in files_order:
        filepath = docs_dir / filename
        if filepath.exists():
            print(f"Traitement de {filename}...")

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            # Ajouter une page de séparation
            doc.add_page_break()

            # Convertir le contenu
            convert_rst_to_docx(content, doc)

    # Sauvegarder le document
    output_path = docs_dir / "Consultator_Documentation.docx"
    doc.save(output_path)

    print(f"Documentation Word créée: {output_path}")
    return output_path


if __name__ == "__main__":
    create_word_documentation()
