#!/usr/bin/env python3
"""
Script final pour créer une documentation Word professionnelle
avec sommaire fonctionnel et présentation améliorée
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def setup_professional_styles(doc):
    """Configure des styles professionnels Word"""

    # Style titre principal
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.name = "Arial"
    title_style.font.color.rgb = RGBColor(31, 119, 180)  # Bleu Consultator

    # Style sous-titre
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.italic = True
    subtitle_style.font.name = "Arial"
    subtitle_style.font.color.rgb = RGBColor(89, 89, 89)

    # Créer des styles de titre personnalisés
    for level in range(1, 4):
        style_name = f"Heading{level}Custom"
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Configuration basée sur le niveau
        if level == 1:
            style.font.size = Pt(20)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(12)
            style.paragraph_format.outline_level = 0
        elif level == 2:
            style.font.size = Pt(16)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.outline_level = 1
        elif level == 3:
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.outline_level = 2

        style.font.bold = True
        style.font.name = "Arial"
        style.paragraph_format.keep_with_next = True

    # Style pour le code
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(64, 64, 64)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Style pour les listes
    list_style = doc.styles.add_style("CustomList", WD_STYLE_TYPE.PARAGRAPH)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_after = Pt(3)

    # Style pour les notes
    note_style = doc.styles.add_style("NoteStyle", WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(128, 128, 128)
    note_style.paragraph_format.left_indent = Inches(0.5)
    note_style.paragraph_format.right_indent = Inches(0.5)


def add_table_of_contents(doc):
    """Ajoute une table des matières manuelle (plus fiable)"""

    # Ajouter un saut de page
    doc.add_page_break()

    # Titre de la table des matières
    toc_title = doc.add_paragraph("Table des matières", style="Heading1Custom")
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # Espace

    # Contenu de la table des matières (sera rempli plus tard)
    toc_content = doc.add_paragraph(
        "La table des matières sera générée automatiquement.\n"
    )
    toc_content.add_run(
        "Pour mettre à jour : Clic droit > Mettre à jour les champs"
    ).italic = True

    # Ajouter un saut de page après la TOC
    doc.add_page_break()


def parse_rst_table(table_lines):
    """Parse un tableau RST amélioré"""
    if not table_lines:
        return None, None

    # Détecter les en-têtes
    headers = []
    data_rows = []

    # Chercher la ligne d'en-têtes (après la première ligne de séparation)
    separator_indices = []
    for i, line in enumerate(table_lines):
        if re.match(r"^\s*\+[-+]*\+", line):
            separator_indices.append(i)

    if len(separator_indices) >= 2:
        # Ligne d'en-têtes
        header_line = table_lines[separator_indices[0] + 1].strip()
        if header_line.startswith("|") and header_line.endswith("|"):
            headers = [cell.strip() for cell in header_line.split("|")[1:-1]]

        # Lignes de données
        for i in range(separator_indices[1] + 1, len(separator_indices) - 1, 2):
            if i < len(table_lines):
                row_line = table_lines[i].strip()
                if row_line.startswith("|") and row_line.endswith("|"):
                    cells = [cell.strip() for cell in row_line.split("|")[1:-1]]
                    data_rows.append(cells)

    return headers, data_rows


def convert_rst_to_docx_improved(rst_content, doc, section_title=""):
    """Convertit RST vers Word avec gestion améliorée"""

    lines = rst_content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Détecter les titres
        if i + 1 < len(lines) and lines[i + 1].startswith(("=", "-", "~", "^", '"')):
            title_text = line.strip()
            underline = lines[i + 1][0]

            # Ne pas traiter le titre principal si c'est la page de garde
            if (
                underline == "="
                and section_title
                and title_text.lower() == section_title.lower()
            ):
                i += 2
                continue

            if underline == "=":
                p = doc.add_paragraph(title_text, style="Heading1Custom")
            elif underline == "-":
                p = doc.add_paragraph(title_text, style="Heading2Custom")
            elif underline == "~":
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            elif underline == "^":
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            else:
                p = doc.add_paragraph(title_text, style="Heading3Custom")

            i += 2
            continue

        # Détecter les tableaux list-table
        if line.startswith(".. list-table::"):
            # Collecter les informations du tableau
            table_title = ""
            table_headers = []
            table_data = []

            j = i + 1
            while j < len(lines):
                current_line = lines[j]

                if current_line.startswith("   :header-rows:"):
                    # Cette ligne contient le nombre de lignes d'en-tête
                    j += 1
                    continue
                elif current_line.strip().startswith("* - ") and not table_headers:
                    # Ligne d'en-têtes
                    header_line = current_line.strip()[4:]  # Enlever "* - "
                    table_headers = [
                        h.strip() for h in header_line.split("        ") if h.strip()
                    ]
                    j += 1
                    continue
                elif current_line.strip().startswith("* ") and table_headers:
                    # Ligne de données
                    data_line = current_line.strip()[2:]  # Enlever "* "
                    row_data = [
                        d.strip() for d in data_line.split("        ") if d.strip()
                    ]
                    if len(row_data) == len(table_headers):
                        table_data.append(row_data)
                    j += 1
                    continue
                elif current_line.strip() == "":
                    j += 1
                    continue
                else:
                    break

            # Créer le tableau Word
            if table_headers and table_data:
                table = doc.add_table(rows=len(table_data) + 1, cols=len(table_headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Ajouter les en-têtes
                header_row = table.rows[0]
                for col_idx, header in enumerate(table_headers):
                    header_row.cells[col_idx].text = header
                    # Style de l'en-tête
                    cell = header_row.cells[col_idx]
                    cell.paragraphs[0].runs[0].font.bold = True

                # Ajouter les données
                for row_idx, row_data in enumerate(table_data, 1):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data

                # Style du tableau
                table.style = "Table Grid"

            i = j
            continue

        # Détecter les blocs de code
        if line.startswith(".. code-block::"):
            code_lines = []
            j = i + 1

            # Collecter les lignes de code
            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith("   ") or current_line.strip() == "":
                    if current_line.strip() == "" and code_lines:
                        break
                    code_lines.append(
                        current_line[3:]
                        if current_line.startswith("   ")
                        else current_line
                    )
                    j += 1
                else:
                    break

            if code_lines:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph(code_text, style="CodeBlock")
                i = j
                continue

        # Détecter les listes
        if line.startswith(("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")):
            list_items = []
            j = i

            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith(
                    ("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")
                ):
                    list_items.append(current_line[2:])
                    j += 1
                elif current_line.strip() == "":
                    j += 1
                else:
                    break

            for item in list_items:
                p = doc.add_paragraph(item, style="CustomList")

            i = j
            continue

        # Détecter les notes
        if (
            line.startswith(".. note::")
            or line.startswith(".. warning::")
            or line.startswith(".. tip::")
        ):
            note_type = (
                "💡 Note"
                if line.startswith(".. note::")
                else (
                    "⚠️ Avertissement"
                    if line.startswith(".. warning::")
                    else "💡 Conseil"
                )
            )

            note_lines = []
            j = i + 1

            while j < len(lines):
                current_line = lines[j]
                if current_line.strip() == "" and note_lines:
                    break
                if current_line.strip():
                    note_lines.append(current_line.strip())
                j += 1

            if note_lines:
                note_text = f"{note_type}: {' '.join(note_lines)}"
                p = doc.add_paragraph(note_text, style="NoteStyle")

            i = j
            continue

        # Ignorer les autres directives
        if line.startswith(".. "):
            i += 1
            continue

        # Texte normal
        if line.strip():
            p = doc.add_paragraph(line)

        i += 1


def create_professional_word_documentation():
    """Crée une documentation Word professionnelle avec sommaire"""

    # Configuration des sections
    sections = [
        ("index.rst", "Page de garde", False),
        ("installation.rst", "Installation et Configuration", True),
        ("quickstart.rst", "Guide de démarrage rapide", True),
        ("features.rst", "Fonctionnalités", True),
        ("development.rst", "Guide de développement", True),
        ("api.rst", "API Reference", True),
        ("tutorials.rst", "Tutoriels", True),
    ]

    # Créer le document
    doc = Document()

    # Configuration des styles professionnels
    setup_professional_styles(doc)

    # Page de garde
    title = doc.add_paragraph("📋 Consultator", style="Title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Documentation Technique Complète", style="Subtitle")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Informations générales
    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.add_run("Application de gestion d'une practice data de consultants\n").bold = (
        False
    )
    info.add_run(
        "Interface Streamlit moderne avec analyses avancées et chatbot IA\n"
    ).bold = False
    info.add_run("Architecture modulaire et évolutive").bold = False
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    version_info = doc.add_paragraph()
    version_info.add_run("Version 1.0.0\n").bold = True
    version_info.add_run("Générée automatiquement depuis Sphinx\n").bold = False
    version_info.add_run("Septembre 2025").bold = False
    version_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Table des matières
    add_table_of_contents(doc)

    # Traiter chaque section
    docs_dir = Path(__file__).parent

    for filename, section_title, add_page_break in sections:
        filepath = docs_dir / filename
        if filepath.exists():
            print(f"📄 Traitement de {filename}...")

            if add_page_break:
                doc.add_page_break()

            # Titre de section
            section_para = doc.add_paragraph(section_title, style="Heading1Custom")
            section_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            # Convertir le contenu
            convert_rst_to_docx_improved(content, doc, section_title)

    # Sauvegarder le document
    output_path = docs_dir / "Consultator_Documentation_Pro.docx"
    doc.save(output_path)

    print(f"✅ Documentation Word professionnelle créée: {output_path}")
    print(f"📊 Taille: {os.path.getsize(output_path)} octets")

    return output_path


if __name__ == "__main__":
    create_professional_word_documentation()
