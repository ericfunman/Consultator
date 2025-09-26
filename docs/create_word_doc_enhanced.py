#!/usr/bin/env python3
"""
Script amélioré pour convertir la documentation Sphinx RST en document Word
Version améliorée avec gestion des tableaux, liens, et mise en forme avancée
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


def setup_document_styles(doc):
    """Configure les styles du document Word avec une meilleure mise en forme"""

    # Style pour les titres principaux
    title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.name = "Arial"
    title_style.font.color.rgb = RGBColor(31, 119, 180)  # Bleu Consultator
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(24)

    # Style pour les sous-titres
    subtitle_style = doc.styles.add_style("CustomSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.italic = True
    subtitle_style.font.name = "Arial"
    subtitle_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_style.paragraph_format.space_after = Pt(36)

    # Style pour les titres de niveau 1
    h1_style = doc.styles.add_style("Heading1Custom", WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.name = "Arial"
    h1_style.font.color.rgb = RGBColor(31, 119, 180)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    h1_style.paragraph_format.keep_with_next = True

    # Style pour les titres de niveau 2
    h2_style = doc.styles.add_style("Heading2Custom", WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.name = "Arial"
    h2_style.font.color.rgb = RGBColor(31, 119, 180)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)
    h2_style.paragraph_format.keep_with_next = True

    # Style pour les titres de niveau 3
    h3_style = doc.styles.add_style("Heading3Custom", WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.name = "Arial"
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # Style pour le code
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(64, 64, 64)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.shading = parse_xml(
        r'<w:shd {} w:fill="F8F8F8"/>'.format(nsdecls("w"))
    )

    # Style pour les listes
    list_style = doc.styles.add_style("CustomList", WD_STYLE_TYPE.PARAGRAPH)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_after = Pt(3)

    # Style pour les notes
    note_style = doc.styles.add_style("CustomNote", WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(128, 128, 128)
    note_style.paragraph_format.left_indent = Inches(0.5)
    note_style.paragraph_format.right_indent = Inches(0.5)
    note_style.paragraph_format.shading = parse_xml(
        r'<w:shd {} w:fill="FFFFE0"/>'.format(nsdecls("w"))
    )


def parse_rst_table(table_lines):
    """Parse un tableau RST et retourne les données"""
    if not table_lines:
        return None

    # Trouver les séparateurs
    separator_indices = []
    for i, line in enumerate(table_lines):
        if re.match(r"^\s*\+[-+]*\+", line):
            separator_indices.append(i)

    if len(separator_indices) < 2:
        return None

    # Extraire les données
    data_rows = []
    for i in range(1, len(separator_indices) - 1, 2):
        row_start = separator_indices[i] + 1
        row_end = separator_indices[i + 1]

        if row_start < len(table_lines):
            row_line = table_lines[row_start].strip()
            if row_line.startswith("|") and row_line.endswith("|"):
                # Diviser par |
                cells = [cell.strip() for cell in row_line.split("|")[1:-1]]
                data_rows.append(cells)

    return data_rows if data_rows else None


def convert_rst_to_docx(rst_content, doc):
    """Convertit le contenu RST en éléments Word avec gestion améliorée"""

    lines = rst_content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Détecter les titres
        if i + 1 < len(lines) and lines[i + 1].startswith(("=", "-", "~", "^", '"')):
            title_text = line.strip()
            underline = lines[i + 1][0]

            if underline == "=":
                # Titre principal
                p = doc.add_paragraph(title_text, style="CustomTitle")
            elif underline == "-":
                # Titre niveau 1
                p = doc.add_paragraph(title_text, style="Heading1Custom")
            elif underline == "~":
                # Titre niveau 2
                p = doc.add_paragraph(title_text, style="Heading2Custom")
            elif underline == "^":
                # Titre niveau 3
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            else:
                # Titre niveau 4 ou plus
                p = doc.add_paragraph(title_text, style="Heading3Custom")

            i += 2  # Sauter la ligne de soulignement
            continue

        # Détecter les tableaux
        if line.startswith(".. list-table::") or (
            line.strip() == "" and i > 0 and lines[i - 1].startswith(".. list-table::")
        ):
            # Collecter les lignes du tableau
            table_lines = []
            j = i

            # Chercher la directive list-table
            while j < len(lines) and not lines[j].startswith(".. list-table::"):
                j += 1

            if j < len(lines):
                # Collecter jusqu'à la fin du tableau
                k = j + 1
                while k < len(lines):
                    current_line = lines[k]
                    if current_line.strip() == "" and k > j + 2:
                        break
                    table_lines.append(current_line)
                    k += 1

                # Parser le tableau
                table_data = parse_rst_table(table_lines)
                if table_data and len(table_data) > 0:
                    # Créer le tableau Word
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_data

                    # Style du tableau
                    table.style = "Table Grid"

                    i = k
                    continue

        # Détecter les blocs de code
        if line.startswith(".. code-block::") or (
            i > 0 and lines[i - 1].startswith(".. code-block::")
        ):
            code_lines = []
            j = i

            # Sauter la directive
            if line.startswith(".. code-block::"):
                j += 1

            # Collecter les lignes de code
            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith("   ") or current_line.strip() == "":
                    if current_line.strip() == "" and code_lines:
                        break
                    code_lines.append(
                        current_line[3:]
                        if current_line.startswith("   ")
                        else current_line
                    )
                    j += 1
                else:
                    break

            if code_lines:
                # Ajouter le bloc de code
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph(code_text, style="CodeBlock")
                i = j
                continue

        # Détecter les listes
        if line.startswith(("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")):
            list_items = []
            j = i

            # Collecter tous les éléments de la liste
            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith(
                    ("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")
                ):
                    list_items.append(current_line[2:])
                    j += 1
                elif current_line.strip() == "":
                    j += 1
                else:
                    break

            # Ajouter les éléments de liste
            for item in list_items:
                p = doc.add_paragraph(item, style="CustomList")

            i = j
            continue

        # Détecter les notes et avertissements
        if (
            line.startswith(".. note::")
            or line.startswith(".. warning::")
            or line.startswith(".. tip::")
        ):
            note_lines = []
            j = i + 1

            # Collecter le contenu de la note
            while j < len(lines):
                current_line = lines[j]
                if current_line.strip() == "" and note_lines:
                    break
                if current_line.strip():
                    note_lines.append(current_line.strip())
                j += 1

            if note_lines:
                note_text = " ".join(note_lines)
                p = doc.add_paragraph(f"💡 {note_text}", style="CustomNote")

            i = j
            continue

        # Détecter les directives à ignorer
        if line.startswith(".. ") and not any(
            line.startswith(f".. {directive}::")
            for directive in ["note", "warning", "tip", "list-table", "code-block"]
        ):
            i += 1
            continue

        # Texte normal
        if line.strip():
            p = doc.add_paragraph(line)

            # Mise en évidence du texte spécial
            if "**" in line:
                # Gras
                pass  # Pourrait être amélioré avec regex
            if "*" in line:
                # Italique
                pass  # Pourrait être amélioré avec regex

        i += 1


def create_enhanced_word_documentation():
    """Crée la documentation Word améliorée"""

    # Fichiers à traiter dans l'ordre
    files_order = [
        ("index.rst", "Page de garde"),
        ("installation.rst", "Installation et Configuration"),
        ("quickstart.rst", "Guide de démarrage rapide"),
        ("features.rst", "Fonctionnalités"),
        ("development.rst", "Guide de développement"),
        ("api.rst", "API Reference"),
        ("tutorials.rst", "Tutoriels"),
    ]

    # Créer le document
    doc = Document()

    # Configuration des styles
    setup_document_styles(doc)

    # Page de garde
    title = doc.add_paragraph("📋 Consultator", style="CustomTitle")
    subtitle = doc.add_paragraph("Documentation Complète", style="CustomSubtitle")

    doc.add_paragraph("")
    doc.add_paragraph("Application de gestion d'une practice data de consultants")
    doc.add_paragraph("Interface Streamlit moderne avec analyses avancées")
    doc.add_paragraph("")
    doc.add_paragraph("Version 1.0.0")
    doc.add_paragraph("Générée automatiquement depuis Sphinx")
    doc.add_paragraph("Septembre 2025")

    # Table des matières
    doc.add_page_break()
    toc_title = doc.add_paragraph("Table des matières", style="Heading1Custom")

    # Traiter chaque fichier
    docs_dir = Path(__file__).parent

    for filename, section_title in files_order:
        filepath = docs_dir / filename
        if filepath.exists():
            print(f"Traitement de {filename}...")

            # Nouvelle section
            doc.add_page_break()
            section_header = doc.add_paragraph(section_title, style="Heading1Custom")

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            # Convertir le contenu (sauter les en-têtes déjà traités)
            lines = content.split("\n")
            # Supprimer les premières lignes si c'est un titre principal
            if (
                lines
                and lines[0].strip()
                and len(lines) > 1
                and lines[1].startswith("=")
            ):
                # Sauter le titre principal car on l'a déjà mis
                start_line = 2
                while start_line < len(lines) and lines[start_line].strip() == "":
                    start_line += 1
                content = "\n".join(lines[start_line:])

            convert_rst_to_docx(content, doc)

    # Sauvegarder le document
    output_path = docs_dir / "Consultator_Documentation_Complete.docx"
    doc.save(output_path)

    print(f"Documentation Word améliorée créée: {output_path}")
    return output_path


if __name__ == "__main__":
    create_enhanced_word_documentation()
