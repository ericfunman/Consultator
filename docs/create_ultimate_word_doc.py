#!/usr/bin/env python3
"""
Script ULTIME pour créer une documentation Word parfaite
avec table des matières cliquable et contenu propre
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def setup_professional_styles(doc):
    """Configure des styles professionnels Word avec niveaux de titre pour TOC"""

    # Style titre principal
    title_style = doc.styles['Title']
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.name = 'Arial'
    title_style.font.color.rgb = RGBColor(31, 119, 180)

    # Style sous-titre
    subtitle_style = doc.styles['Subtitle']
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.italic = True
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.color.rgb = RGBColor(89, 89, 89)

    # Créer des styles de titre avec outline levels pour TOC automatique
    for level in range(1, 4):
        style_name = f'Heading{level}Custom'
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        if level == 1:
            style.font.size = Pt(20)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(12)
            style.paragraph_format.outline_level = 0  # Niveau 1 dans TOC
        elif level == 2:
            style.font.size = Pt(16)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.outline_level = 1  # Niveau 2 dans TOC
        elif level == 3:
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.outline_level = 2  # Niveau 3 dans TOC

        style.font.bold = True
        style.font.name = 'Arial'
        style.paragraph_format.keep_with_next = True

    # Style pour le code
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(64, 64, 64)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Style pour les listes
    list_style = doc.styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_after = Pt(3)

    # Style pour les notes
    note_style = doc.styles.add_style('NoteStyle', WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(128, 128, 128)
    note_style.paragraph_format.left_indent = Inches(0.5)
    note_style.paragraph_format.right_indent = Inches(0.5)

def extract_toc_from_rst(content):
    """Extrait proprement la structure TOC depuis RST"""
    toc_sections = []
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line == '.. toctree::':
            toctree_info = {'maxdepth': 2, 'caption': '', 'items': []}

            # Parser les options
            j = i + 1
            while j < len(lines):
                option_line = lines[j].strip()

                if option_line.startswith(':maxdepth:'):
                    try:
                        toctree_info['maxdepth'] = int(option_line.split(':')[2].strip())
                    except:
                        pass
                elif option_line.startswith(':caption:'):
                    # EXTRAIRE UNIQUEMENT LE TEXTE APRÈS ":caption: "
                    caption_full = option_line.split(':', 1)[1].strip()
                    toctree_info['caption'] = caption_full
                elif option_line == '':
                    j += 1
                    continue
                elif not option_line.startswith(':'):
                    break
                j += 1

            # Collecter les items
            while j < len(lines):
                item_line = lines[j].strip()
                if item_line and not item_line.startswith('..') and not item_line.startswith(':'):
                    clean_item = item_line.replace('.rst', '').replace('/', ' > ')
                    toctree_info['items'].append(clean_item)
                elif item_line == '':
                    j += 1
                    continue
                else:
                    break
                j += 1

            if toctree_info['items']:
                toc_sections.append(toctree_info)

            i = j
        else:
            i += 1

    return toc_sections

def add_clickable_table_of_contents(doc, toc_sections):
    """Ajoute une table des matières avec outline levels pour TOC automatique"""

    # Ajouter un saut de page
    doc.add_page_break()

    # Titre de la table des matières
    toc_title = doc.add_paragraph("Table des matières", style='Heading1Custom')
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # Espace

    # Générer la structure avec les bons outline levels
    for section in toc_sections:
        if section['caption']:
            # Titre de section comme Heading2Custom (outline level 1)
            section_title = doc.add_paragraph(section['caption'], style='Heading2Custom')

            # Items comme Heading3Custom (outline level 2)
            for item in section['items']:
                item_para = doc.add_paragraph(f"• {item}", style='Heading3Custom')

    # Instructions pour la TOC automatique
    doc.add_paragraph("")
    toc_instructions = doc.add_paragraph("📋 Table des matières automatique :", style='Heading2Custom')
    toc_instructions.add_run("\n1. Allez dans Références > Table des matières")
    toc_instructions.add_run("\n2. Choisissez un style automatique")
    toc_instructions.add_run("\n3. La TOC sera générée avec des liens cliquables")

    # Saut de page après TOC
    doc.add_page_break()

def should_skip_line(line, in_grid_section=False):
    """Filtre COMPLET des lignes RST à ignorer"""
    line = line.strip()

    # Images et options d'images
    if line.startswith('.. image::') or line.startswith('   :alt:') or line.startswith('   :align:') or line.startswith('   :width:'):
        return True

    # Éléments de mise en page
    if line in ['|', '.. grid::', '   :gutter:']:
        return True

    # Grid items et leurs options
    if line.startswith('   .. grid-item-card::') or line.startswith('      :link:') or line.startswith('      :link-type:'):
        return True

    # Contenu des grid-item-card (lignes suivantes)
    if in_grid_section and (line.startswith('      ') and not line.startswith('      :')):
        return True

    # Tableaux de métriques (éviter qu'ils aillent dans TOC)
    if line.startswith('.. list-table::') and ('Métriques' in line or 'Ressources' in line):
        return True

    # Options de tableaux
    if line.startswith('   :header-rows:') or line.startswith('   :widths:'):
        return True

    # Lignes de données de tableau
    if line.startswith('* ') and (' - ' in line or '     ' in line):
        return True

    # Liens |today| et autres substitutions
    if line.strip() in ['|today|', '| **Version:** 1.0.0']:
        return True

    return False

def convert_rst_to_docx_perfect(rst_content, doc, section_title="", is_index=False):
    """Conversion RST vers Word parfaite avec filtrage ultime"""

    lines = rst_content.split('\n')
    i = 0

    skip_toctree = False
    in_grid_section = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Détecter le début d'une section grid
        if line.strip() == '.. grid::':
            in_grid_section = True
            i += 1
            continue

        # Détecter la fin d'une section grid (ligne vide après)
        if in_grid_section and line.strip() == '':
            in_grid_section = False

        # Filtrer les lignes à ignorer
        if should_skip_line(line, in_grid_section):
            i += 1
            continue

        # Ignorer toctree
        if line.strip() == '.. toctree::':
            skip_toctree = True
            i += 1
            continue
        elif skip_toctree:
            if line.strip() == '' or line.strip().startswith(':') or not line.strip().startswith('..'):
                i += 1
                continue
            else:
                skip_toctree = False
                continue

        # Détecter les titres
        if i + 1 < len(lines) and lines[i + 1].startswith(('=', '-', '~', '^', '"')):
            title_text = line.strip()
            underline = lines[i + 1][0]

            # Ne pas traiter le titre principal si c'est index
            if underline == '=' and is_index and 'Documentation Complète' in title_text:
                i += 2
                continue

            if underline == '=':
                p = doc.add_paragraph(title_text, style='Heading1Custom')
            elif underline == '-':
                p = doc.add_paragraph(title_text, style='Heading2Custom')
            elif underline == '~':
                p = doc.add_paragraph(title_text, style='Heading3Custom')
            elif underline == '^':
                p = doc.add_paragraph(title_text, style='Heading3Custom')
            else:
                p = doc.add_paragraph(title_text, style='Heading3Custom')

            i += 2
            continue

        # Détecter les blocs de code
        if line.startswith('.. code-block::'):
            code_lines = []
            j = i + 1

            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith('   ') or current_line.strip() == '':
                    if current_line.strip() == '' and code_lines:
                        break
                    code_lines.append(current_line[3:] if current_line.startswith('   ') else current_line)
                    j += 1
                else:
                    break

            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text, style='CodeBlock')
                i = j
                continue

        # Détecter les listes
        if line.startswith(('- ', '* ', '+ ', '1. ', '2. ', '3. ', '• ')):
            list_items = []
            j = i

            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith(('- ', '* ', '+ ', '1. ', '2. ', '3. ', '• ')):
                    list_items.append(current_line[2:])
                    j += 1
                elif current_line.strip() == '':
                    j += 1
                else:
                    break

            for item in list_items:
                p = doc.add_paragraph(item, style='CustomList')

            i = j
            continue

        # Détecter les notes
        if line.startswith('.. note::') or line.startswith('.. warning::') or line.startswith('.. tip::'):
            note_type = "💡 Note" if line.startswith('.. note::') else "⚠️ Avertissement" if line.startswith('.. warning::') else "💡 Conseil"

            note_lines = []
            j = i + 1

            while j < len(lines):
                current_line = lines[j]
                if current_line.strip() == '' and note_lines:
                    break
                if current_line.strip():
                    note_lines.append(current_line.strip())
                j += 1

            if note_lines:
                note_text = f"{note_type}: {' '.join(note_lines)}"
                p = doc.add_paragraph(note_text, style='NoteStyle')

            i = j
            continue

        # Tableaux list-table (mais pas les métriques)
        if line.startswith('.. list-table::') and 'Métriques' not in line and 'Ressources' not in line:
            # Logique de tableau simplifiée
            i += 1
            continue

        # Ignorer les autres directives
        if line.startswith('.. '):
            i += 1
            continue

        # Texte normal
        if line.strip():
            p = doc.add_paragraph(line)

        i += 1

def create_ultimate_word_documentation():
    """Crée la documentation Word ultime avec TOC cliquable"""

    # Créer le document
    doc = Document()

    # Styles professionnels
    setup_professional_styles(doc)

    # Importer les modules nécessaires pour la TOC
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml

    # Lire index.rst pour extraire la TOC
    docs_dir = Path(__file__).parent
    index_path = docs_dir / 'index.rst'

    toc_sections = []
    if index_path.exists():
        with open(index_path, 'r', encoding='utf-8') as f:
            index_content = f.read()
        toc_sections = extract_toc_from_rst(index_content)

    # Page de garde
    title = doc.add_paragraph("📋 Consultator", style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Documentation Technique Complète", style='Subtitle')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.add_run("Application de gestion d'une practice data de 60 consultants\n").bold = False
    info.add_run("Interface Streamlit moderne avec analyses avancées et chatbot IA\n").bold = False
    info.add_run("Architecture modulaire et évolutive").bold = False
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    version_info = doc.add_paragraph()
    version_info.add_run("Version 1.0.0\n").bold = True
    version_info.add_run("Générée automatiquement depuis Sphinx\n").bold = False
    version_info.add_run("Septembre 2025").bold = False
    version_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Table des matières cliquable
    add_clickable_table_of_contents(doc, toc_sections)

    # Traiter chaque section
    sections_processed = set()

    # Index (page d'accueil)
    if index_path.exists():
        print(f"📄 Traitement de index.rst...")
        doc.add_paragraph("Page d'accueil", style='Heading1Custom')

        with open(index_path, 'r', encoding='utf-8') as f:
            content = f.read()

        convert_rst_to_docx_perfect(content, doc, "Page d'accueil", is_index=True)
        sections_processed.add('index')

    # Autres sections
    for section in toc_sections:
        for item in section['items']:
            filename = item.split(' > ')[0]
            if filename in sections_processed:
                continue

            filepath = docs_dir / f"{filename}.rst"
            if filepath.exists():
                print(f"📄 Traitement de {filename}.rst...")
                doc.add_page_break()

                section_title = section['caption']
                section_para = doc.add_paragraph(section_title, style='Heading1Custom')

                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()

                convert_rst_to_docx_perfect(content, doc, section_title)

                sections_processed.add(filename)

    # Sauvegarder
    output_path = docs_dir / "Consultator_Documentation_Ultime.docx"
    doc.save(output_path)

    print(f"✅ Documentation Word ultime créée: {output_path}")
    print(f"📊 Taille: {os.path.getsize(output_path)} octets")
    print(f"📋 Sections dans la TOC: {len(toc_sections)}")

    return output_path

if __name__ == "__main__":
    create_ultimate_word_documentation()
