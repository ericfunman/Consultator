#!/usr/bin/env python3
"""
Script final pour créer une documentation Word professionnelle
avec sommaire parfait et filtrage complet du contenu RST
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


def setup_professional_styles(doc):
    """Configure des styles professionnels Word"""

    # Style titre principal
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.name = "Arial"
    title_style.font.color.rgb = RGBColor(31, 119, 180)  # Bleu Consultator

    # Style sous-titre
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.italic = True
    subtitle_style.font.name = "Arial"
    subtitle_style.font.color.rgb = RGBColor(89, 89, 89)

    # Créer des styles de titre personnalisés
    for level in range(1, 4):
        style_name = f"Heading{level}Custom"
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Configuration basée sur le niveau
        if level == 1:
            style.font.size = Pt(20)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(12)
            style.paragraph_format.outline_level = 0
        elif level == 2:
            style.font.size = Pt(16)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.outline_level = 1
        elif level == 3:
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(31, 119, 180)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.outline_level = 2

        style.font.bold = True
        style.font.name = "Arial"
        style.paragraph_format.keep_with_next = True

    # Style pour le code
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(64, 64, 64)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Style pour les listes
    list_style = doc.styles.add_style("CustomList", WD_STYLE_TYPE.PARAGRAPH)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_after = Pt(3)

    # Style pour les notes
    note_style = doc.styles.add_style("NoteStyle", WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(128, 128, 128)
    note_style.paragraph_format.left_indent = Inches(0.5)
    note_style.paragraph_format.right_indent = Inches(0.5)

    # Style pour les liens
    link_style = doc.styles.add_style("LinkStyle", WD_STYLE_TYPE.PARAGRAPH)
    link_style.font.color.rgb = RGBColor(0, 102, 204)
    link_style.font.underline = True


def extract_toc_from_rst(content):
    """Extrait la structure de la table des matières depuis le RST"""
    toc_sections = []
    lines = content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Détecter une directive toctree
        if line == ".. toctree::":
            toctree_info = {"maxdepth": 2, "caption": "", "items": []}

            # Parser les options de la directive
            j = i + 1
            while j < len(lines):
                option_line = lines[j].strip()

                if option_line.startswith(":maxdepth:"):
                    try:
                        toctree_info["maxdepth"] = int(
                            option_line.split(":")[2].strip()
                        )
                    except:
                        pass
                elif option_line.startswith(":caption:"):
                    # Nettoyer la caption (enlever "caption: " du début)
                    caption_text = option_line.split(":", 1)[1].strip()
                    toctree_info["caption"] = caption_text
                elif option_line == "":
                    j += 1
                    continue
                elif not option_line.startswith(":"):
                    # C'est le début des items
                    break
                j += 1

            # Collecter les items
            while j < len(lines):
                item_line = lines[j].strip()
                if (
                    item_line
                    and not item_line.startswith("..")
                    and not item_line.startswith(":")
                ):
                    # Nettoyer le nom de l'item (enlever .rst et remplacer / par >)
                    clean_item = item_line.replace(".rst", "").replace("/", " > ")
                    toctree_info["items"].append(clean_item)
                elif item_line == "":
                    j += 1
                    continue
                else:
                    break
                j += 1

            if toctree_info["items"]:
                toc_sections.append(toctree_info)

            i = j
        else:
            i += 1

    return toc_sections


def add_table_of_contents(doc, toc_sections):
    """Ajoute une table des matières basée sur les sections toctree"""

    # Ajouter un saut de page
    doc.add_page_break()

    # Titre de la table des matières
    toc_title = doc.add_paragraph("Table des matières", style="Heading1Custom")
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # Espace

    # Générer le contenu de la table des matières
    for section in toc_sections:
        if section["caption"]:
            # Titre de section (sans "caption: ")
            section_title = doc.add_paragraph(
                section["caption"], style="Heading2Custom"
            )

            # Items de la section
            for item in section["items"]:
                item_para = doc.add_paragraph(f"• {item}", style="CustomList")

    # Note sur la mise à jour
    doc.add_paragraph("")
    update_note = doc.add_paragraph(
        "💡 Pour mettre à jour automatiquement : Clic droit sur la table des matières > Mettre à jour les champs",
        style="NoteStyle",
    )

    # Ajouter un saut de page après la TOC
    doc.add_page_break()


def should_skip_line(line):
    """Détermine si une ligne doit être ignorée dans la conversion"""
    line = line.strip()

    # Ignorer les directives d'image
    if (
        line.startswith(".. image::")
        or line.startswith("   :alt:")
        or line.startswith("   :align:")
        or line.startswith("   :width:")
    ):
        return True

    # Ignorer les séparateurs | et autres éléments de mise en page RST
    if line in [
        "|",
        ".. grid::",
        "   :gutter:",
        "   .. grid-item-card::",
        "      :link:",
        "      :link-type:",
    ]:
        return True

    # Ignorer les options de grid-item-card
    if (
        line.startswith("   .. grid-item-card::")
        or line.startswith("      :link:")
        or line.startswith("      :link-type:")
    ):
        return True

    # Ignorer les directives de métriques
    if line.startswith(".. list-table::") and "Métriques" in line:
        return True

    # Ignorer les options de tableau
    if line.startswith("   :header-rows:") or line.startswith("   :widths:"):
        return True

    # Ignorer les lignes de données de tableau qui commencent par *
    if line.startswith("* ") and (" - " in line or "     " in line):
        return True

    return False


def convert_rst_to_docx_improved(
    rst_content, doc, section_title="", toc_sections=None, is_index=False
):
    """Convertit RST vers Word avec filtrage complet du contenu RST"""

    lines = rst_content.split("\n")
    i = 0

    # Variables pour gérer l'état
    skip_toctree = False
    in_table = False
    table_headers = []
    table_data = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Ignorer les lignes qui doivent être filtrées
        if should_skip_line(line):
            i += 1
            continue

        # Détecter les directives toctree et les ignorer (déjà traité dans la TOC)
        if line.strip() == ".. toctree::":
            skip_toctree = True
            i += 1
            continue
        elif skip_toctree:
            if (
                line.strip() == ""
                or line.strip().startswith(":")
                or not line.strip().startswith("..")
            ):
                i += 1
                continue
            else:
                skip_toctree = False
                continue

        # Détecter les titres
        if i + 1 < len(lines) and lines[i + 1].startswith(("=", "-", "~", "^", '"')):
            title_text = line.strip()
            underline = lines[i + 1][0]

            # Ne pas traiter le titre principal si c'est la page de garde
            if (
                underline == "="
                and section_title
                and title_text.lower() == section_title.lower()
            ):
                i += 2
                continue

            if underline == "=":
                p = doc.add_paragraph(title_text, style="Heading1Custom")
            elif underline == "-":
                p = doc.add_paragraph(title_text, style="Heading2Custom")
            elif underline == "~":
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            elif underline == "^":
                p = doc.add_paragraph(title_text, style="Heading3Custom")
            else:
                p = doc.add_paragraph(title_text, style="Heading3Custom")

            i += 2
            continue

        # Détecter les blocs de code
        if line.startswith(".. code-block::"):
            code_lines = []
            j = i + 1

            # Collecter les lignes de code
            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith("   ") or current_line.strip() == "":
                    if current_line.strip() == "" and code_lines:
                        break
                    code_lines.append(
                        current_line[3:]
                        if current_line.startswith("   ")
                        else current_line
                    )
                    j += 1
                else:
                    break

            if code_lines:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph(code_text, style="CodeBlock")
                i = j
                continue

        # Détecter les listes
        if line.startswith(("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")):
            list_items = []
            j = i

            while j < len(lines):
                current_line = lines[j]
                if current_line.startswith(
                    ("- ", "* ", "+ ", "1. ", "2. ", "3. ", "• ")
                ):
                    list_items.append(current_line[2:])
                    j += 1
                elif current_line.strip() == "":
                    j += 1
                else:
                    break

            for item in list_items:
                p = doc.add_paragraph(item, style="CustomList")

            i = j
            continue

        # Détecter les notes
        if (
            line.startswith(".. note::")
            or line.startswith(".. warning::")
            or line.startswith(".. tip::")
        ):
            note_type = (
                "💡 Note"
                if line.startswith(".. note::")
                else (
                    "⚠️ Avertissement"
                    if line.startswith(".. warning::")
                    else "💡 Conseil"
                )
            )

            note_lines = []
            j = i + 1

            while j < len(lines):
                current_line = lines[j]
                if current_line.strip() == "" and note_lines:
                    break
                if current_line.strip():
                    note_lines.append(current_line.strip())
                j += 1

            if note_lines:
                note_text = f"{note_type}: {' '.join(note_lines)}"
                p = doc.add_paragraph(note_text, style="NoteStyle")

            i = j
            continue

        # Ignorer les autres directives
        if line.startswith(".. "):
            i += 1
            continue

        # Texte normal
        if line.strip():
            p = doc.add_paragraph(line)

        i += 1


def create_professional_word_documentation():
    """Crée une documentation Word professionnelle avec sommaire parfait"""

    # Créer le document
    doc = Document()

    # Configuration des styles professionnels
    setup_professional_styles(doc)

    # Lire le fichier index.rst pour extraire la structure TOC
    docs_dir = Path(__file__).parent
    index_path = docs_dir / "index.rst"

    toc_sections = []
    if index_path.exists():
        with open(index_path, "r", encoding="utf-8") as f:
            index_content = f.read()
        toc_sections = extract_toc_from_rst(index_content)

    # Page de garde
    title = doc.add_paragraph("📋 Consultator", style="Title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Documentation Technique Complète", style="Subtitle")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Informations générales
    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.add_run(
        "Application de gestion d'une practice data de 60 consultants\n"
    ).bold = False
    info.add_run(
        "Interface Streamlit moderne avec analyses avancées et chatbot IA\n"
    ).bold = False
    info.add_run("Architecture modulaire et évolutive").bold = False
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    version_info = doc.add_paragraph()
    version_info.add_run("Version 1.0.0\n").bold = True
    version_info.add_run("Générée automatiquement depuis Sphinx\n").bold = False
    version_info.add_run("Septembre 2025").bold = False
    version_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Table des matières avec contenu réel et captions nettoyées
    add_table_of_contents(doc, toc_sections)

    # Traiter chaque section
    sections_processed = set()

    # D'abord traiter index.rst (mais filtrer le toctree et autres éléments)
    if index_path.exists():
        print(f"📄 Traitement de index.rst...")
        doc.add_paragraph("Page d'accueil", style="Heading1Custom")

        with open(index_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Convertir le contenu en filtrant tout le contenu RST moche
        convert_rst_to_docx_improved(
            content, doc, "Page d'accueil", toc_sections, is_index=True
        )
        sections_processed.add("index")

    # Traiter les autres sections selon la structure TOC
    for section in toc_sections:
        for item in section["items"]:
            # Nettoyer le nom du fichier
            filename = item.split(" > ")[0]  # Prendre la première partie avant >
            if filename in sections_processed:
                continue

            filepath = docs_dir / f"{filename}.rst"
            if filepath.exists():
                print(f"📄 Traitement de {filename}.rst...")

                doc.add_page_break()

                # Titre de section basé sur la caption
                section_title = section["caption"]
                section_para = doc.add_paragraph(section_title, style="Heading1Custom")
                section_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()

                # Convertir le contenu
                convert_rst_to_docx_improved(content, doc, section_title)

                sections_processed.add(filename)

    # Sauvegarder le document
    output_path = docs_dir / "Consultator_Documentation_Final.docx"
    doc.save(output_path)

    print(f"✅ Documentation Word professionnelle créée: {output_path}")
    print(f"📊 Taille: {os.path.getsize(output_path)} octets")
    print(f"📋 Sections dans la table des matières: {len(toc_sections)}")

    return output_path


if __name__ == "__main__":
    create_professional_word_documentation()
